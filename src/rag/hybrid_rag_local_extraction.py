# pip install weaviate-client rdflib datasets tqdm python-dotenv requests sentence-transformers torch PyMuPDF python-docx markitdown docling langextract
# Refs: https://github.com/WoSea/AI-Document-Assistant/blob/main/src/backend/ingest/extract_text.py
'''
Convert various document formats to markdown text.
convert_file_to_markdown(path, prefer_order=...):

markitdown (preferred - converts to Markdown) - best: produces Markdown directly
PyMuPDF  for PDF => text => produce minimal markdown - extract PDF text page-by-page
docling for DOCX  - document parsing library
langextract used for language extraction/cleanup
python-docx as robust DOCX fallback
Plain-text fallback (read .txt)

When upserting to Weaviate, the source_doc contains the original filename and chunk indices; text contains the markdown chunk.
Put PDF/DOCX files in "./data_files/" or pass a list of files to ingest_files(...)
'''

"""
 Hybrid RAG with Weaviate + RDF + Ollama (Mistral), now with:
  - PDF/DOCX => Markdown conversion (using markitdown / PyMuPDF / docling / docx fallback)
  - Chunking of Markdown text
  - Upsert chunks into Weaviate (server-side Ollama vectorizer / client-side vectors)
  - Hybrid search and RAG generation as before
"""

import os
import sys
import json
import numpy as np
import time
import subprocess
import tempfile
import requests
from typing import List, Tuple, Dict, Any
import shutil

from datasets import load_dataset
from rdflib import Graph, Literal, Namespace, RDF
from tqdm import tqdm

# Weaviate client
import weaviate
from weaviate.classes.config import Configure
from weaviate.classes.config import Property, DataType

# Embedding fallback (local)
from sentence_transformers import SentenceTransformer

# Evaluation metrics
from evaluation_for_rag import recall_at_k, precision_at_k, f1_at_k, reciprocal_rank, compute_generation_metrics, pipeline_metrics

# ---------- optional libs (may or may not be installed) ----------
try:
    import pymupdf   # PyMuPDF
    _HAS_PyMuPDF = True
except Exception:
    _HAS_PyMuPDF = False

try:
    import markitdown  # Microsoft markitdown
    _HAS_MARKITDOWN = True
except Exception:
    markitdown = None
    _HAS_MARKITDOWN = False

try:
    import docling
    _HAS_DOCLING = True
except Exception:
    docling = None
    _HAS_DOCLING = False

try:
    import langextract
    _HAS_LANGEXTRACT = True
except Exception:
    langextract = None
    _HAS_LANGEXTRACT = False

try:
    import docx  # python-docx fallback
    _HAS_PYDOCX = True
except Exception:
    docx = None
    _HAS_PYDOCX = False

# ========== CONFIG ==========
WEAVIATE_URL = os.getenv("WEAVIATE_URL", "http://localhost:8080")
WEAVIATE_CLASS = "DocChunk"
WEAVIATE_VECTORIZE_WITH_OLLAMA = True
WEAVIATE_OLLAMA_ENDPOINT = os.getenv("OLLAMA_API", "http://host.docker.internal:11434")
WEAVIATE_OLLAMA_EMBED_MODEL = "nomic-embed-text"

USE_OLLAMA_LOCAL_FOR_LLM = True
OLLAMA_LLM_MODEL = "mistral"
USE_MISTRAL_API = False
MISTRAL_API_KEY = os.getenv("MISTRAL_API_KEY", "")

EMBED_LOCAL_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
CHUNK_MAX_WORDS = 120
CHUNK_OVERLAP_WORDS = 20
TOP_K = 3
HYBRID_ALPHA = 0.5

# ======= chunking helper (same as before) =======
def chunk_text_by_words(text: str, max_words: int = CHUNK_MAX_WORDS, overlap: int = CHUNK_OVERLAP_WORDS
                       ) -> List[Tuple[str, int, int]]:
    """
    Split text into chunks by words with a sliding window and overlap.
    Returns list of (chunk_text, start_word_idx, end_word_idx)
    """
    words = text.split()
    if not words:
        return []
    chunks = []
    i = 0
    n = len(words)
    while i < n:
        start = i
        end = min(i + max_words, n)
        chunk = " ".join(words[start:end])
        chunks.append((chunk, start, end))
        if end == n:
            break
        i = end - overlap
    return chunks

# ========= File extraction & Markdown conversion =========
def convert_with_markitdown(input_path: str) -> str:
    """
    Try to convert a document to Markdown using the 'markitdown' CLI if available.
    This calls the markitdown CLI via subprocess (common install method).
    Returns markdown text as string.
    """
    # Prefer CLI because the Python API for markitdown may not be installed
    try:
        out_md = None
        with tempfile.NamedTemporaryFile(suffix=".md", delete=False) as tmp:
            tmp_path = tmp.name
        # markitdown CLI usage (assumes markitdown is on PATH)
        cmd = ["markitdown", input_path, "-o", tmp_path]
        subprocess.run(cmd, check=True)
        with open(tmp_path, "r", encoding="utf-8") as f:
            out_md = f.read()
        os.remove(tmp_path)
        return out_md
    except Exception as e:
        # If CLI not available, try Python package if installed (best-effort)
        if _HAS_MARKITDOWN:
            try:
                # If markitdown python bindings exist (best-effort)
                if hasattr(markitdown, "markdown_from_path"):
                    return markitdown.markdown_from_path(input_path)
                elif hasattr(markitdown, "Convert"):
                    c = markitdown.Convert(input_path)
                    return c.to_markdown()
            except Exception:
                pass
        # fallback - raise so caller can try other extractors
        raise e

def extract_text_from_pdf_pymupdf(path: str) -> str:
    """
    Use PyMuPDF to extract text page-by-page and produce minimal Markdown (page headings).
    """
    if not _HAS_PyMuPDF:
        raise RuntimeError("PyMuPDF is not installed")
    doc = pymupdf.open(path) # open a document
    pages = []
    i = 0
    for page in doc: # iterate the document pages
        text = page.get_text() # get plain text encoded as UTF-8
        pages.append(f"## Page {i+1}\n\n{text.strip()}\n")
        i += 1
    doc.close()
    return "\n".join(pages)

def extract_text_from_docx_python_docx(path: str) -> str:
    """
    Extract paragraphs from a .docx using python-docx and return simple Markdown (paragraphs as blank-line separated).
    """
    if not _HAS_PYDOCX:
        raise RuntimeError("python-docx is not installed")
    doc = docx.Document(path)
    paras = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paras.append(text)
    # join paragraphs with blank lines -> basic markdown
    return "\n\n".join(paras)

def convert_file_to_markdown(filepath: str, prefer_order: List[str] = None) -> str:
    """
    Convert PDF/DOCX/TXT -> Markdown using available tools.
    prefer_order: list of method names in order of preference, e.g. ["markitdown","pymupdf","docling","docx"]
    The function will try each and return the first successful conversion result.
    """
    if prefer_order is None:
        prefer_order = ["markitdown", "pymupdf", "docling", "docx", "txt"]

    ext = os.path.splitext(filepath)[1].lower()
    last_exc = None

    for method in prefer_order:
        try:
            if method == "markitdown":
                # try markitdown (best: produces Markdown)
                if _HAS_MARKITDOWN or shutil.which("markitdown"):
                    md = convert_with_markitdown(filepath)
                    print(f"[extract] used markitdown for {filepath}")
                    return md
                else:
                    continue

            if method == "pymupdf" and ext == ".pdf":
                if _HAS_PyMuPDF:
                    md = extract_text_from_pdf_pymupdf(filepath)
                    print(f"[extract] used PyMuPDF for {filepath}")
                    return md
                else:
                    continue

            if method == "docling" and ext in (".docx", ".doc", ".pdf"):
                if _HAS_DOCLING:
                    # docling usage may vary; attempt a generic call if available
                    try:
                        # docling may expose a Reader to load documents
                        reader = getattr(docling, "Reader", None)
                        if reader:
                            r = reader(filepath)
                            text = r.text if hasattr(r, "text") else "\n".join(r.paragraphs)
                            print(f"[extract] used docling.Reader for {filepath}")
                            return text
                        # fallback to docling.parse_file
                        if hasattr(docling, "parse_file"):
                            text = docling.parse_file(filepath)
                            print(f"[extract] used docling.parse_file for {filepath}")
                            return text
                    except Exception as e:
                        last_exc = e
                        continue
                else:
                    continue

            if method == "docx" and ext in (".docx",):
                if _HAS_PYDOCX:
                    md = extract_text_from_docx_python_docx(filepath)
                    print(f"[extract] used python-docx for {filepath}")
                    return md
                else:
                    continue

            if method == "txt" and ext in (".txt",):
                # simple read
                with open(filepath, "r", encoding="utf-8") as f:
                    txt = f.read()
                print(f"[extract] used plain text read for {filepath}")
                return txt

            # Add additional handlers if needed (e.g. call langextract tools)
        except Exception as e:
            last_exc = e
            continue

    # If reached here, nothing succeeded; raise informative error
    raise RuntimeError(f"Failed to convert {filepath} to markdown. Last error: {last_exc}")

# ========== embedding helpers (same as before) ==========
_local_embedder = None
def init_local_embedder():
    global _local_embedder
    if _local_embedder is None:
        _local_embedder = SentenceTransformer(EMBED_LOCAL_MODEL)
    return _local_embedder

def embed_texts_local(texts: List[str]) -> List[List[float]]:
    model = init_local_embedder()
    embs = model.encode(texts, normalize_embeddings=True)
    return embs.tolist()

def embed_with_ollama_api(texts: List[str]) -> List[List[float]]:
    out = []
    url = f"{WEAVIATE_OLLAMA_ENDPOINT.rstrip('/')}/api/embeddings"
    for t in texts:
        payload = {"model": WEAVIATE_OLLAMA_EMBED_MODEL, "prompt": t}
        r = requests.post(url, json=payload, timeout=60)
        r.raise_for_status()
        j = r.json()
        vec = j.get("embedding") or j.get("data", [{}])[0].get("embedding")
        out.append(vec)
    return out

def embed_texts(texts: List[str]) -> List[List[float]]:
    try:
        return embed_with_ollama_api(texts)
    except Exception:
        return embed_texts_local(texts)

# =======---- Weaviate helpers (per new quickstart) =======----
def connect_weaviate_client() -> weaviate.WeaviateClient:
    client = weaviate.connect_to_local(url=WEAVIATE_URL)
    return client

def create_weaviate_collection(client: weaviate.WeaviateClient, class_name: str):
    try:
        if client.collections.exists(class_name):
            client.collections.delete(class_name)
            time.sleep(0.5)
    except Exception:
        pass

    if WEAVIATE_VECTORIZE_WITH_OLLAMA:
        vec_conf = Configure.Vectors.text2vec_ollama(
            api_endpoint=WEAVIATE_OLLAMA_ENDPOINT,
            model=WEAVIATE_OLLAMA_EMBED_MODEL
        )
        gen_conf = Configure.Generative.ollama(
            api_endpoint=WEAVIATE_OLLAMA_ENDPOINT,
            model=OLLAMA_LLM_MODEL
        )
        client.collections.create(
            name=class_name,
            vector_config=vec_conf,
            generative_config=gen_conf,
            properties=[
                Property(name="text", data_type=DataType.TEXT),
                Property(name="source_doc", data_type=DataType.TEXT),
                Property(name="chunk_start", data_type=DataType.INT),
                Property(name="chunk_end", data_type=DataType.INT),
                Property(name="category", data_type=DataType.TEXT),
                Property(name="region", data_type=DataType.TEXT),
            ],
        )
    else:
        client.collections.create(
            name=class_name,
            vector_config=Configure.Vectors.none(),
            properties=[
                Property(name="text", data_type=DataType.TEXT),
                Property(name="source_doc", data_type=DataType.TEXT),
                Property(name="chunk_start", data_type=DataType.INT),
                Property(name="chunk_end", data_type=DataType.INT),
                Property(name="category", data_type=DataType.TEXT),
                Property(name="region", data_type=DataType.TEXT),
            ],
        )

def upsert_chunks_to_weaviate(client: weaviate.WeaviateClient, class_name: str,
                              chunks: List[Dict[str, Any]], vectors: List[List[float]] = None):
    col = client.collections.use(class_name)
    with col.batch.fixed_size(batch_size=200) as batch:
        for i, chunk in enumerate(chunks):
            vec = None
            if vectors:
                vec = vectors[i]
            batch.add_object(
                properties={
                    "text": chunk["text"],
                    "source_doc": chunk["source_doc"],
                    "chunk_start": int(chunk["start"]),
                    "chunk_end": int(chunk["end"]),
                    "category": chunk.get("category", "Healthcare"),
                    "region": chunk.get("region", "US"),
                },
                vector=vec
            )
            if batch.number_errors > 10:
                print("Too many batch errors, abort.")
                break

# =========== retrieval / rag helpers (same as before) ===========
def weaviate_hybrid_search(client: weaviate.WeaviateClient, class_name: str,
                           query_text: str, query_vector: List[float] = None,
                           top_k: int = TOP_K, alpha: float = HYBRID_ALPHA, filters: Any = None):
    col = client.collections.use(class_name)
    if query_vector is not None:
        res = col.query.hybrid(query=query_text, vector=query_vector, alpha=alpha, limit=top_k,
                               return_properties=["text", "source_doc", "chunk_start", "chunk_end", "category", "region"])
    else:
        res = col.query.near_text(query=query_text, limit=top_k,
                                  return_properties=["text", "source_doc", "chunk_start", "chunk_end", "category", "region"])
    hits = [o.properties.get("text") for o in res.objects]
    return hits

def weaviate_generate_with_context(client: weaviate.WeaviateClient, class_name: str, query_text: str, top_k: int = TOP_K):
    col = client.collections.use(class_name)
    resp = col.generate.near_text(query=query_text, limit=top_k, grouped_task="Answer using retrieved facts.")
    return getattr(resp.generative, "text", "")

# =========== LLM direct calls ===========
def llm_generate_with_ollama(prompt: str) -> str:
    url = f"{WEAVIATE_OLLAMA_ENDPOINT.rstrip('/')}/api/generate"
    payload = {"model": OLLAMA_LLM_MODEL, "prompt": prompt, "stream": False}
    r = requests.post(url, json=payload, timeout=60)
    r.raise_for_status()
    j = r.json()
    return j.get("response", "") or j.get("text", "")

def llm_generate_with_mistral_api(prompt: str) -> str:
    url = "https://api.mistral.ai/v1/chat/completions"
    headers = {"Authorization": f"Bearer {MISTRAL_API_KEY}", "Content-Type": "application/json"}
    payload = {"model": "open-mistral-7b", "messages": [{"role": "user", "content": prompt}], "temperature": 0.2}
    r = requests.post(url, headers=headers, json=payload, timeout=60)
    r.raise_for_status()
    j = r.json()
    return j["choices"][0]["message"]["content"]

def llm_generate_direct(prompt: str) -> str:
    if USE_OLLAMA_LOCAL_FOR_LLM:
        return llm_generate_with_ollama(prompt)
    elif USE_MISTRAL_API:
        return llm_generate_with_mistral_api(prompt)
    else:
        raise RuntimeError("No LLM configured.")

# =========== Main flow: file ingest + chunking + weaviate ingest ===========
def ingest_files(file_paths: List[str], prefer_order: List[str] = None, limit_files: int = 50):
    """
    Convert files -> markdown -> chunk -> return list of chunks.
    prefer_order: list like ["markitdown","pymupdf","docling","docx","txt"]
    """
    chunks = []
    file_count = 0
    for fp in file_paths:
        if file_count >= limit_files:
            break
        try:
            md = convert_file_to_markdown(fp, prefer_order=prefer_order)
        except Exception as e:
            print(f"[ingest] Skipping {fp}: extraction failed: {e}")
            continue
        # optional: minimal metadata - basename
        base = os.path.basename(fp)
        # chunk markdown text
        doc_chunks = chunk_text_by_words(md, max_words=CHUNK_MAX_WORDS, overlap=CHUNK_OVERLAP_WORDS)
        for chunk_text, start, end in doc_chunks:
            chunks.append({
                "text": chunk_text,
                "source_doc": base,
                "start": start,
                "end": end,
                "category": "Healthcare",
                "region": "US"
            })
        file_count += 1
    return chunks

# ====== Scoring helpers ======
def cosine_similarity(vec_a: List[float], vec_b: List[float]) -> float:
    a = np.array(vec_a)
    b = np.array(vec_b)
    return float(np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b) + 1e-8))

# ====== Reciprocal Rank Fusion (RRF) ======
def reciprocal_rank_fusion(results_per_query: List[List[str]], k: int = 60, top_k: int = TOP_K) -> List[Tuple[str, float]]:
    """
    results_per_query: list of results list, each list is [chunk_text,...] for 1 query
    Returns a list [(chunk_text, fused_score)] sorted in descending order
    """
    scores = {}
    for result_list in results_per_query:
        for rank, doc in enumerate(result_list):
            score = 1.0 / (k + rank + 1)   # +1 to avoid division by zero
            scores[doc] = scores.get(doc, 0.0) + score
    ranked = sorted(scores.items(), key=lambda x: x[1], reverse=True)
    return ranked[:top_k]

# ========== Query Rewrite module ==========
def rewrite_query_with_llm(query: str, n: int = 3) -> List[str]:
    """
    Using LLM (Ollama/Mistral) to generate rewritten versions of the query.
    Returns a list of queries: [original query + paraphrases].
    """
    prompt = f"""
You are a query rewriting assistant.
Rewrite the following question into {n} alternative queries
that preserve the original meaning but use different words.
Return them as a numbered list.

Original query: "{query}"
"""
    try:
        rewritten = llm_generate_direct(prompt)
        # parse numbered list
        lines = [line.strip("0123456789). ") for line in rewritten.split("\n") if line.strip()]
        rewrites = [q for q in lines if len(q.split()) > 2]
        return [query] + rewrites[:n]
    except Exception as e:
        print("Query rewrite failed:", e)
        return [query]

# ========== Query rewrite + hybrid search + combine results ==========
def weaviate_multiquery_search_old(client: weaviate.WeaviateClient,
                               class_name: str,
                               query_text: str,
                               top_k: int = TOP_K,
                               alpha: float = HYBRID_ALPHA) -> List[str]:
    queries = rewrite_query_with_llm(query_text, n=3)
    seen = set()
    results = []

    for q in queries:
        try:
            q_vec = embed_texts([q])[0]
        except Exception:
            q_vec = None
        hits = weaviate_hybrid_search(client, class_name, q, query_vector=q_vec,
                                      top_k=top_k, alpha=alpha)
        for h in hits:
            if h not in seen:
                seen.add(h)
                results.append(h)
    return results[:top_k]

# Query rewrite + Hybrid search for multiple queries + Merge results + Semantic re-ranking
def weaviate_multiquery_search(client: weaviate.WeaviateClient,
                               class_name: str,
                               query_text: str,
                               top_k: int = TOP_K,
                               alpha: float = HYBRID_ALPHA) -> List[Tuple[str, float]]:
    queries = rewrite_query_with_llm(query_text, n=3)
    seen = {}
    results = []

    # embed user query to ranking
    q_vec_main = embed_texts([query_text])[0]

    for q in queries:
        try:
            q_vec = embed_texts([q])[0]
        except Exception:
            q_vec = None

        hits = weaviate_hybrid_search(
            client, class_name, q, query_vector=q_vec,
            top_k=top_k, alpha=alpha
        )

        for h in hits:
            if h not in seen:
                seen[h] = True
                # embed chunk text to calculate similarity
                try:
                    h_vec = embed_texts([h])[0]
                    score = cosine_similarity(q_vec_main, h_vec)
                except Exception:
                    score = 0.0
                results.append((h, score))

    # re-rank by score
    results.sort(key=lambda x: x[1], reverse=True)
    return results[:top_k]

# ====== Multi-query search by RRF ====== Query rewrite + Hybrid search + Reciprocal Rank Fusion (RRF) for merging and re-rank results
def weaviate_multiquery_rrf_search(client: weaviate.WeaviateClient,
                                   class_name: str,
                                   query_text: str,
                                   top_k: int = TOP_K,
                                   alpha: float = HYBRID_ALPHA) -> List[Tuple[str, float]]:
    queries = rewrite_query_with_llm(query_text, n=3)
    results_per_query = []

    for q in queries:
        try:
            q_vec = embed_texts([q])[0]
        except Exception:
            q_vec = None

        hits = weaviate_hybrid_search(
            client, class_name, q,
            query_vector=q_vec, top_k=top_k, alpha=alpha
        )
        results_per_query.append(hits)

    fused_results = reciprocal_rank_fusion(results_per_query, k=60, top_k=top_k)
    return fused_results

# ===== Evaluation Loop =====
def evaluate_pipeline(client, class_name: str, queries: List[Dict[str, Any]],
                      search_fn, embed_fn, k: int = 5):
    results = {
        "Recall@K": [], "Precision@K": [], "F1@K": [], "MRR": [],
        "BLEU": [], "ROUGE-L": [], "METEOR": [], "FactualAcc": [],
        "Latency": [], "TokenCost": []
    }
    answers_all, refs_all = [], []

    for q in tqdm(queries, desc="Evaluating"):
        query_text = q["query"]
        relevant_docs = q.get("relevant_docs", [])
        gt_answers = q.get("answers", [])

        # Retrieval
        start = time.time()
        ranked = search_fn(client, class_name, query_text, top_k=k)
        latency = time.time() - start

        retrieved = [t for t, _ in ranked]

        # Retrieval metrics
        results["Recall@K"].append(recall_at_k(retrieved, relevant_docs, k))
        results["Precision@K"].append(precision_at_k(retrieved, relevant_docs, k))
        results["F1@K"].append(f1_at_k(retrieved, relevant_docs, k))
        results["MRR"].append(reciprocal_rank(retrieved, relevant_docs))

        # Generation
        context = "\n".join(retrieved[:k])
        prompt = f"Answer concisely using the context:\n{context}\n\nQ: {query_text}\nA:"
        answer = llm_generate_direct(prompt)
        g_metrics = compute_generation_metrics(answer, gt_answers)

        results["BLEU"].append(g_metrics["BLEU"])
        results["ROUGE-L"].append(g_metrics["ROUGE-L"])
        results["METEOR"].append(g_metrics["METEOR"])
        results["FactualAcc"].append(g_metrics["FactualAcc"])

        # Pipeline
        results["Latency"].append(latency)
        results["TokenCost"].append(len(answer.split()))
        answers_all.append(answer)
        refs_all.append(gt_answers)

    # Aggregate retrieval + generation
    metrics = {m: float(np.mean(vals)) for m, vals in results.items()}

    # Add pipeline metrics
    pipeline_m = pipeline_metrics(results["Latency"], results["TokenCost"], answers_all, refs_all)
    metrics.update(pipeline_m)

    return metrics

# ========== Script entrypoint ==========
def main():
    # find files in ./data_files (pdf/docx/txt)
    src_dir = "src\\data\\data_files"
    if not os.path.exists(src_dir):
        print(f"Create '{src_dir}' and drop your .pdf/.docx/.txt files then run again.")
        return

    file_paths = []
    for root, _, files in os.walk(src_dir):
        for f in files:
            if f.lower().endswith((".pdf", ".docx", ".txt")):
                file_paths.append(os.path.join(root, f))

    if not file_paths:
        print("No files found in data_files/ (pdf/docx/txt). Exiting.")
        return

    print(f"Found {len(file_paths)} files to ingest (showing first 50):")
    for fp in file_paths[:50]:
        print(" -", fp)

    # ingest -> chunks
    chunks = ingest_files(file_paths, prefer_order=["markitdown","pymupdf","docling","docx","txt"], limit_files=50)
    print(f"Total chunks prepared: {len(chunks)}")

    # connect weaviate and create class
    client = connect_weaviate_client()
    print("Weaviate ready:", client.is_ready())
    create_weaviate_collection(client, WEAVIATE_CLASS)
    time.sleep(0.3)

    # if are not letting Weaviate vectorize, compute vectors locally
    vectors = None
    if not WEAVIATE_VECTORIZE_WITH_OLLAMA:
        print("Computing local embeddings for chunks...")
        texts = [c["text"] for c in chunks]
        vectors = embed_texts(texts)
        print("Computed vectors for chunks.")

    # upsert batches
    print("Upserting chunks into Weaviate...")
    upsert_chunks_to_weaviate(client, WEAVIATE_CLASS, chunks, vectors=vectors)
    print("Upsert done.")

    # sample query
    user_query = "What healthcare benefits are mentioned across these documents?"
    print("Preparing query embedding (client-side) ...")
    try:
        q_vec = embed_texts([user_query])[0]
    except Exception as e:
        print("Query embed failed, will rely on server-side:", e)
        q_vec = None

    # semantic retrieval + ranking
    # # # top_docs = weaviate_hybrid_search(client, WEAVIATE_CLASS, user_query, query_vector=q_vec, top_k=TOP_K, alpha=HYBRID_ALPHA)
    # # top_docs = weaviate_multiquery_search_old(client, WEAVIATE_CLASS, user_query, top_k=TOP_K, alpha=HYBRID_ALPHA)

    # # print("Top retrieved snippets:")
    # # for i, t in enumerate(top_docs, 1):
    # #     print(f"{i}. {t[:300]}...")

    # ranked_docs = weaviate_multiquery_search(client, WEAVIATE_CLASS, user_query, top_k=TOP_K, alpha=HYBRID_ALPHA)
    # print("Top retrieved snippets (re-ranked):")
    # top_docs = []
    # for i, (t, score) in enumerate(ranked_docs, 1):
    #     print(f"{i}. [score={score:.4f}] {t[:300]}...")
    #     top_docs.append(t)

    ranked_docs = weaviate_multiquery_rrf_search(client, WEAVIATE_CLASS, user_query, top_k=TOP_K, alpha=HYBRID_ALPHA)
    print("Top retrieved snippets (RRF fused):")
    top_docs = []
    for i, (t, score) in enumerate(ranked_docs, 1):
        print(f"{i}. [RRF={score:.4f}] {t[:300]}...")
        top_docs.append(t)
    
    # Compose prompt and call LLM
    context = "\n".join(top_docs[:TOP_K])
    prompt = f"Answer using the context below. Provide a concise summary.\n\n{context}\n\nQ: {user_query}\nA:"
    print("Calling LLM for final answer...")
    answer = llm_generate_direct(prompt)
    print("\n=== FINAL ANSWER ===")
    print(answer)

    # Evaluate pipeline
    eval_queries = [
        {"query": "What healthcare benefits are mentioned?", "relevant": ["health insurance", "medical benefits"]},
        {"query": "What is the leave policy?", "relevant": ["annual leave", "paid time off"]},
    ]
    print("\nRunning evaluation pipeline...")
    metrics = evaluate_pipeline(client, WEAVIATE_CLASS, eval_queries, top_k=TOP_K)
    print("=== Evaluation Metrics ===")
    for k, v in metrics.items():
        print(f"{k}: {v:.4f}")

    client.close()

if __name__ == "__main__":
    main()